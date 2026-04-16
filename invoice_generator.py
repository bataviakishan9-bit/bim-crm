"""
invoice_generator.py — Generate Word (.docx) and PDF invoices
matching the BIM INFRASOLUTIONS LLP template.
"""
import io
from datetime import datetime
from num2words import num2words as _n2w


# ── helpers ───────────────────────────────────────────────────────────────────

def _amount_words(amount: float) -> str:
    """Convert amount to words (e.g. 17700 → 'Seventeen thousand seven hundred only')."""
    try:
        rupees = int(amount)
        paise  = round((amount - rupees) * 100)
        words  = _n2w(rupees, lang="en_IN").capitalize()
        if paise:
            paise_w = _n2w(paise, lang="en_IN").capitalize()
            return f"{words} rupees and {paise_w} paise only"
        return f"{words} only"
    except Exception:
        return f"INR {amount:.2f} only"


def _fmt_date(date_str: str) -> str:
    """Convert YYYY-MM-DD to DD.MM.YYYY."""
    try:
        return datetime.strptime(date_str[:10], "%Y-%m-%d").strftime("%d.%m.%Y")
    except Exception:
        return date_str


# ── WORD generation ───────────────────────────────────────────────────────────

def generate_docx(invoice: dict, items: list) -> bytes:
    """Return .docx bytes for the given invoice."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    def _para(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold      = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def _set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)

    # ── Header: Company name ──────────────────────────────────────────────────
    h = doc.add_heading("BIM INFRASOLUTIONS LLP", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(27, 58, 107)

    # GSTIN / LLPIN row
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.style = doc.styles["Normal Table"]
    t.cell(0, 0).text = f"GSTIN: {invoice.get('company_gstin','24AAUFB9689E1ZS')}"
    t.cell(0, 1).text = f"LLPIN: {invoice.get('llpin','AAP-1096')}"
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # ── Tax Invoice title ─────────────────────────────────────────────────────
    _para("Tax Invoice", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # Date + Invoice no
    date_str = _fmt_date(invoice.get("date", ""))
    inv_no   = invoice.get("invoice_no", "")
    _para(f"Date: {date_str}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(f"Invoice no: {inv_no}", size=10)

    lut = invoice.get("lut_number", "")
    if lut:
        _para(f"LUT Number: {lut}", size=9)

    doc.add_paragraph()

    # ── To section ────────────────────────────────────────────────────────────
    _para("To,", bold=True, size=10)
    _para(invoice.get("client_name", ""), bold=True, size=10)
    addr = invoice.get("client_address", "")
    if addr:
        _para(addr, size=10)
    cgstin = invoice.get("client_gstin", "")
    if cgstin:
        _para(f"GSTIN: {cgstin}", bold=True, size=10)

    doc.add_paragraph()

    # ── Items table ───────────────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for cell, text in zip(hdr, ["Particulars", "Unit", "Amount (INR)"]):
        _cell_text(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255,255,255))
        _set_cell_bg(cell, "1B3A6B")

    subtotal = 0.0
    for item in items:
        row  = tbl.add_row().cells
        desc = item.get("description", "")
        sac  = item.get("sac_code", "")
        if sac:
            desc += f" (SAC: {sac})"
        amt  = float(item.get("amount", 0))
        subtotal += amt
        _cell_text(row[0], desc, size=10)
        _cell_text(row[1], str(item.get("unit", 1)), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row[2], f"{amt:,.0f}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    gst_rate   = float(invoice.get("gst_rate", 18))
    gst_amount = float(invoice.get("gst_amount", 0))
    total      = float(invoice.get("total", 0))

    # GST row
    gst_row = tbl.add_row().cells
    _cell_text(gst_row[0], f"GST {int(gst_rate)}%", size=10)
    _cell_text(gst_row[1], "", size=10)
    _cell_text(gst_row[2], f"{gst_amount:,.0f}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Total row
    tot_row = tbl.add_row().cells
    _cell_text(tot_row[0], "Total Bill", bold=True, size=10)
    _cell_text(tot_row[1], "-", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(tot_row[2], f"{total:,.0f} INR", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in tot_row:
        _set_cell_bg(cell, "EBF5FB")

    # Amount in words row
    words_row = tbl.add_row().cells
    merged = words_row[0].merge(words_row[1]).merge(words_row[2])
    _cell_text(merged, f"Total amount in words: {_amount_words(total)}", bold=True, size=10,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_bg(merged, "EBF5FB")

    # Column widths
    from docx.shared import Cm
    widths = [Cm(10), Cm(2.5), Cm(4)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    doc.add_paragraph()

    # ── Notes ─────────────────────────────────────────────────────────────────
    notes = invoice.get("notes", "")
    if notes:
        _para(f"Notes: {notes}", size=9, space_after=6)

    # ── Signature ─────────────────────────────────────────────────────────────
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig_p.add_run("Mr. Kishan Batavia\n\nAuthorized Signatory")
    sig_run.bold = True
    sig_run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Terms ─────────────────────────────────────────────────────────────────
    _para("Terms and Conditions:", bold=True, size=10)
    terms = [
        "Subject to Ahmedabad Jurisdiction.",
        "The payment shall be released within two weeks from the date of invoice.",
        "In case of electronic fund transfer banking information is as under:",
    ]
    for i, t in enumerate(terms, 1):
        _para(f"{i}. {t}", size=9)

    doc.add_paragraph()

    # ── Bank details table ────────────────────────────────────────────────────
    _para("ACCOUNT DETAILS", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    bank = [
        ("Beneficiary account name", "BIM INFRASOLUTIONS LLP"),
        ("PAN",                       "AAUFB9689E"),
        ("IFSC code",                 "HDFC0003905"),
        ("Account no",                "50200041261501"),
        ("Account branch",            "ST XAVIERS COLLEGE ROAD"),
        ("Branch code",               "3905"),
        ("Address",                   "HDFC Bank Ltd, Shop No. 7/8, Gr Floor, Zodiac Plaza, Navrangpura, Ahmedabad – 380009, Gujarat, India."),
        ("Swift code",                "HDFCINBB"),
        ("GST no",                    "24AAUFB9689E1ZS"),
    ]
    btbl = doc.add_table(rows=len(bank), cols=2)
    btbl.style = "Table Grid"
    for i, (k, v) in enumerate(bank):
        _cell_text(btbl.rows[i].cells[0], k, bold=True, size=9)
        _cell_text(btbl.rows[i].cells[1], v, size=9)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _para(
        "302, Shahibaug Greens, Shahibaug, Ahmedabad, Gujarat, India. "
        "Phone: +91 8511184053, +64 225053122  |  "
        "E-Mail: info@biminfrasolutions.in  |  Website: www.biminfrasolutions.in",
        size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF generation ────────────────────────────────────────────────────────────

def generate_pdf(invoice: dict, items: list) -> bytes:
    """Return PDF bytes using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                     Paragraph, Spacer, HRFlowable)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=1.5*cm, bottomMargin=1.5*cm)

    styles = getSampleStyleSheet()
    BIM_BLUE = colors.HexColor("#1B3A6B")
    LIGHT_BLUE = colors.HexColor("#EBF5FB")

    def S(name, **kw):
        base = styles[name]
        return ParagraphStyle(name + "_custom", parent=base, **kw)

    story = []

    # Company header
    story.append(Paragraph(
        "<b>BIM INFRASOLUTIONS LLP</b>",
        S("Title", fontSize=20, textColor=BIM_BLUE, alignment=TA_CENTER, spaceAfter=2)
    ))

    # GSTIN / LLPIN
    story.append(Table(
        [[f"GSTIN: {invoice.get('company_gstin','24AAUFB9689E1ZS')}",
          f"LLPIN: {invoice.get('llpin','AAP-1096')}"]],
        colWidths=["50%", "50%"],
        style=[
            ("FONTSIZE",    (0,0), (-1,-1), 8),
            ("ALIGN",       (0,0), (0,0),   "LEFT"),
            ("ALIGN",       (1,0), (1,0),   "RIGHT"),
            ("BOTTOMPADDING",(0,0),(-1,-1), 4),
        ]
    ))
    story.append(HRFlowable(width="100%", thickness=2, color=BIM_BLUE, spaceAfter=6))

    # Tax Invoice title
    story.append(Paragraph("<b>Tax Invoice</b>",
                            S("Normal", fontSize=14, alignment=TA_CENTER, spaceAfter=4)))

    # Date & Invoice no
    date_str = _fmt_date(invoice.get("date", ""))
    inv_no   = invoice.get("invoice_no", "")
    story.append(Paragraph(f"<b>Date:</b> {date_str}",
                            S("Normal", fontSize=10, alignment=TA_RIGHT, spaceAfter=2)))
    story.append(Paragraph(f"<b>Invoice no:</b> {inv_no}",
                            S("Normal", fontSize=10, spaceAfter=2)))
    lut = invoice.get("lut_number", "")
    if lut:
        story.append(Paragraph(f"LUT Number: {lut}",
                                S("Normal", fontSize=9, spaceAfter=2)))

    story.append(Spacer(1, 0.3*cm))

    # To section
    story.append(Paragraph("To,", S("Normal", fontSize=10, fontName="Helvetica-Bold")))
    story.append(Paragraph(f"<b>{invoice.get('client_name','')}</b>",
                            S("Normal", fontSize=10, spaceAfter=1)))
    addr = invoice.get("client_address", "")
    if addr:
        story.append(Paragraph(addr, S("Normal", fontSize=10, spaceAfter=1)))
    cgstin = invoice.get("client_gstin", "")
    if cgstin:
        story.append(Paragraph(f"<b>GSTIN: {cgstin}</b>",
                                S("Normal", fontSize=10, spaceAfter=4)))

    story.append(Spacer(1, 0.4*cm))

    # Items table
    tbl_data = [["Particulars", "Unit", "Amount (INR)"]]
    subtotal = 0.0
    for item in items:
        desc = item.get("description", "")
        sac  = item.get("sac_code", "")
        if sac:
            desc += f"\n(SAC: {sac})"
        amt  = float(item.get("amount", 0))
        subtotal += amt
        tbl_data.append([desc, str(item.get("unit", 1)), f"{amt:,.0f}"])

    gst_rate   = float(invoice.get("gst_rate", 18))
    gst_amount = float(invoice.get("gst_amount", 0))
    total      = float(invoice.get("total", 0))

    tbl_data.append([f"GST {int(gst_rate)}%", "", f"{gst_amount:,.0f}"])
    tbl_data.append(["Total Bill", "-", f"{total:,.0f} INR"])
    tbl_data.append([Paragraph(f"<b>Total amount in words: {_amount_words(total)}</b>",
                                S("Normal", fontSize=9)), "", ""])

    n = len(tbl_data)
    tbl_style = TableStyle([
        # Header row
        ("BACKGROUND",    (0,0),  (-1,0),  BIM_BLUE),
        ("TEXTCOLOR",     (0,0),  (-1,0),  colors.white),
        ("FONTNAME",      (0,0),  (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0,0),  (-1,-1), 9),
        ("ALIGN",         (1,0),  (1,-1),  "CENTER"),
        ("ALIGN",         (2,0),  (2,-1),  "RIGHT"),
        ("GRID",          (0,0),  (-1,-2), 0.5, colors.grey),
        ("BOX",           (0,0),  (-1,-1), 0.5, colors.grey),
        # Total rows highlight
        ("BACKGROUND",    (0,n-3), (-1,n-1), LIGHT_BLUE),
        ("FONTNAME",      (0,n-3), (-1,n-1), "Helvetica-Bold"),
        # Words row spans all cols
        ("SPAN",          (0,n-1), (-1,n-1)),
        ("ALIGN",         (0,n-1), (-1,n-1), "CENTER"),
        ("TOPPADDING",    (0,0),  (-1,-1), 4),
        ("BOTTOMPADDING", (0,0),  (-1,-1), 4),
    ])

    story.append(Table(tbl_data,
                       colWidths=[10*cm, 2.5*cm, 4*cm],
                       style=tbl_style,
                       repeatRows=1))

    story.append(Spacer(1, 0.5*cm))

    # Notes
    notes = invoice.get("notes", "")
    if notes:
        story.append(Paragraph(f"Notes: {notes}", S("Normal", fontSize=9, spaceAfter=4)))

    # Signature
    story.append(Paragraph(
        "<b>Mr. Kishan Batavia</b><br/><br/><b>Authorized Signatory</b>",
        S("Normal", fontSize=10, alignment=TA_RIGHT, spaceAfter=8)
    ))

    # Terms
    story.append(Paragraph("<b>Terms and Conditions:</b>", S("Normal", fontSize=10, spaceAfter=2)))
    for i, t in enumerate([
        "Subject to Ahmedabad Jurisdiction.",
        "The payment shall be released within two weeks from the date of invoice.",
        "In case of electronic fund transfer banking information is as under:",
    ], 1):
        story.append(Paragraph(f"{i}. {t}", S("Normal", fontSize=9, spaceAfter=1)))

    story.append(Spacer(1, 0.3*cm))

    # Bank details
    story.append(Paragraph("<b>ACCOUNT DETAILS</b>",
                            S("Normal", fontSize=10, alignment=TA_CENTER, spaceAfter=4)))
    bank = [
        ["Beneficiary account name", "BIM INFRASOLUTIONS LLP"],
        ["PAN",                       "AAUFB9689E"],
        ["IFSC code",                 "HDFC0003905"],
        ["Account no",                "50200041261501"],
        ["Account branch",            "ST XAVIERS COLLEGE ROAD"],
        ["Branch code",               "3905"],
        ["Address",                   "HDFC Bank Ltd, Shop No. 7/8, Gr Floor, Zodiac Plaza,\nNavrangpura, Ahmedabad – 380009, Gujarat, India."],
        ["Swift code",                "HDFCINBB"],
        ["GST no",                    "24AAUFB9689E1ZS"],
    ]
    story.append(Table(
        [[Paragraph(f"<b>{k}</b>", S("Normal", fontSize=9)),
          Paragraph(v, S("Normal", fontSize=9))] for k, v in bank],
        colWidths=[5*cm, 11.5*cm],
        style=[
            ("GRID",          (0,0), (-1,-1), 0.3, colors.grey),
            ("TOPPADDING",    (0,0), (-1,-1), 3),
            ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ]
    ))

    # Footer
    story.append(HRFlowable(width="100%", thickness=1, color=BIM_BLUE, spaceBefore=8, spaceAfter=4))
    story.append(Paragraph(
        "302, Shahibaug Greens, Shahibaug, Ahmedabad, Gujarat, India. "
        "Phone: +91 8511184053, +64 225053122  |  "
        "info@biminfrasolutions.in  |  www.biminfrasolutions.in",
        S("Normal", fontSize=7.5, alignment=TA_CENTER)
    ))

    doc.build(story)
    return buf.getvalue()
